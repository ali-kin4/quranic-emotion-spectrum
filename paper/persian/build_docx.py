"""Build manuscript.docx from manuscript.md with proper Persian RTL styling.

Pipeline:
  1. Preprocess manuscript.md → strip LaTeX-only commands, convert \emph{}/\textbf{}
     to markdown italic/bold, drop xepersian-specific wrappers.
  2. Build a styled reference.docx (Vazirmatn 12pt body RTL, Amiri for Arabic,
     EB Garamond for Latin, proper heading hierarchy).
  3. Run pandoc with --reference-doc to convert .md → .docx.
  4. Post-process the .docx via python-docx: set body direction RTL, embed
     fonts, fix table styles, set page A4 + margins matching the LaTeX template.

Usage:
    python build_docx.py
"""
from __future__ import annotations

import re
import subprocess
import sys
import zipfile
from pathlib import Path

HERE = Path(__file__).resolve().parent
ROOT = HERE.parent.parent

MD_IN = HERE / "manuscript.md"
MD_TMP = HERE / "_manuscript_for_docx.md"
REF_DOCX = HERE / "reference.docx"
OUT_DOCX = HERE / "manuscript.docx"


# ===================================================================
# 1) Preprocess markdown
# ===================================================================
def preprocess_markdown(text: str) -> str:
    """Strip LaTeX-only commands; convert LaTeX text-formatting to markdown."""
    # Replace the englishabstract environment with a styled English block.
    # We turn it into a plain markdown block and add a marker so post-processing
    # can apply LTR direction + EB Garamond font.
    def _strip_eng(m):
        body = m.group(1)
        # Convert \textbf{...} → **...**, \textit{...} → *...*
        body = re.sub(r'\\textbf\{([^}]+)\}', r'**\1**', body)
        body = re.sub(r'\\textit\{([^}]+)\}', r'*\1*', body)
        body = re.sub(r'\\emph\{([^}]+)\}', r'*\1*', body)
        # Strip LR wrappers (Word handles direction natively)
        body = re.sub(r'\\lr\{([^}]+)\}', r'\1', body)
        # Strip layout-only LaTeX
        body = re.sub(r'\\begin\{center\}', '', body)
        body = re.sub(r'\\end\{center\}', '', body)
        body = re.sub(r'\\noindent\b', '', body)
        body = re.sub(r'\\itshape\b', '', body)
        body = re.sub(r'\\normalfont\b', '', body)
        body = re.sub(r'\\vspace\{[^}]+\}', '', body)
        body = re.sub(r'\\par\b', '', body)
        body = re.sub(r'\\\\\s*', '\n\n', body)
        body = re.sub(r'\{[^{}]*\\bfseries[^{}]*\}', '', body)
        body = re.sub(r'\\Large|\\large|\\small|\\bfseries', '', body)
        body = re.sub(r'\\setstretch\{[^}]+\}', '', body)
        return "\n\n<!--ENGLISH_BLOCK_START-->\n\n" + body + "\n\n<!--ENGLISH_BLOCK_END-->\n\n"

    text = re.sub(
        r'\\begin\{englishabstract\}(.*?)\\end\{englishabstract\}',
        _strip_eng, text, flags=re.DOTALL,
    )

    # Replace LTR/begingroup blocks (bibliography wrapper) with marker
    text = re.sub(r'\\begin\{LTR\}\s*\n', "\n<!--LTR_BLOCK_START-->\n", text)
    text = re.sub(r'\\end\{LTR\}', "\n<!--LTR_BLOCK_END-->\n", text)
    text = re.sub(r'\\englishbody\b', '', text)
    text = re.sub(r'\\addfontfeatures\{[^}]+\}', '', text)
    text = re.sub(r'\\begingroup\b', '', text)
    text = re.sub(r'\\endgroup\b', '', text)
    text = re.sub(r'\\setlength\{[^}]+\}\{[^}]+\}', '', text)
    text = re.sub(r'\\everypar\{[^}]+\}', '', text)
    text = re.sub(r'\\noindent\b', '', text)
    text = re.sub(r'\\itshape\b', '', text)
    text = re.sub(r'\\normalfont\b', '', text)

    # Strip per-line \lr{...} wraps (Word handles direction natively)
    text = re.sub(r'\\lr\{([^{}]+)\}', r'\1', text)

    # Convert other LaTeX formatting to markdown
    text = re.sub(r'\\textbf\{([^}]+)\}', r'**\1**', text)
    text = re.sub(r'\\textit\{([^}]+)\}', r'*\1*', text)
    text = re.sub(r'\\emph\{([^}]+)\}', r'*\1*', text)

    # Drop \arb{...} arabic wrapper, keep content
    text = re.sub(r'\\arb\{([^}]+)\}', r'\1', text)

    # Inline LR text (keep content)
    text = re.sub(r'\\textlatin\{([^}]+)\}', r'\1', text)

    # Figures: swap .pdf -> .png for max Word compatibility.
    # Markdown image syntax: ![caption](file.pdf){width=N%}
    text = re.sub(r'\]\(([^)]+)\.pdf\)\{', r'](\1.png){', text)

    return text


# ===================================================================
# 2) Build the reference.docx with proper styles
# ===================================================================
def build_reference_docx():
    """Modify pandoc's default reference.docx to use Persian RTL fonts."""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    base = HERE / "reference_base.docx"
    if not base.exists():
        subprocess.check_call(
            ["pandoc", "--print-default-data-file=reference.docx",
             "-o", str(base)], cwd=str(HERE))
    # pandoc --print-default-data-file actually writes to stdout, but cmd already
    # captured to base via shell redirect — re-extract just in case
    if not base.exists():
        out = subprocess.check_output(
            ["pandoc", "--print-default-data-file", "reference.docx"])
        base.write_bytes(out)

    doc = Document(str(base))

    # Page setup: A4, 2.6cm margins (match template.tex)
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2.6)
        section.bottom_margin = Cm(2.6)
        section.left_margin = Cm(2.6)
        section.right_margin = Cm(2.6)

    def set_rtl(paragraph_format):
        """Add RTL bidi direction to a paragraph format properties element."""
        # python-docx doesn't expose bidi directly; manipulate XML
        pPr = paragraph_format.element
        bidi = pPr.find(qn('w:bidi'))
        if bidi is None:
            bidi = OxmlElement('w:bidi')
            pPr.append(bidi)

    def set_font(run_or_style, *, latin="EB Garamond", complex_script="Vazirmatn",
                 size_pt=12, bold=False, italic=False):
        """Apply font names + size to a run or style. Sets complex-script font
        (used by Word for RTL/Arabic/Hebrew) separately from Latin."""
        rPr = run_or_style.element
        # rFonts element
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            # Find rPr child; if rPr is style, use first child
            if rPr.tag.endswith('}style'):
                # Look for or create rPr in style
                style_rPr = rPr.find(qn('w:rPr'))
                if style_rPr is None:
                    style_rPr = OxmlElement('w:rPr')
                    rPr.insert(0, style_rPr)
                rPr = style_rPr
                existing = rPr.find(qn('w:rFonts'))
                if existing is not None:
                    rFonts = existing
                else:
                    rPr.insert(0, rFonts)
            else:
                rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), latin)
        rFonts.set(qn('w:hAnsi'), latin)
        rFonts.set(qn('w:cs'), complex_script)
        rFonts.set(qn('w:eastAsia'), latin)

    # --- Body / Normal style: Vazirmatn 12pt, RTL, justified ---
    normal = doc.styles['Normal']
    normal.font.name = "EB Garamond"
    normal.font.size = Pt(12)
    pf = normal.paragraph_format
    pf.line_spacing = 1.35
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.first_line_indent = Cm(0.5)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Force RTL on the body
    pPr = normal.element.get_or_add_pPr()
    bidi = pPr.find(qn('w:bidi'))
    if bidi is None:
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)
    # Set complex-script (RTL) font to Vazirmatn at the style level
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), 'EB Garamond')
    rFonts.set(qn('w:hAnsi'), 'EB Garamond')
    rFonts.set(qn('w:cs'), 'Vazirmatn')
    # Complex-script font size
    szCs = rPr.find(qn('w:szCs'))
    if szCs is None:
        szCs = OxmlElement('w:szCs')
        rPr.append(szCs)
    szCs.set(qn('w:val'), '24')  # 24 half-points = 12pt

    # --- Heading styles ---
    heading_specs = [
        ('Heading 1', 17, True),
        ('Heading 2', 14, True),
        ('Heading 3', 12.5, True),
        ('Heading 4', 11.5, True),
    ]
    for style_name, size, bold in heading_specs:
        try:
            st = doc.styles[style_name]
        except KeyError:
            continue
        st.font.name = "Vazirmatn"
        st.font.size = Pt(size)
        st.font.bold = bold
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        st.paragraph_format.space_before = Pt(10)
        st.paragraph_format.space_after = Pt(4)
        pPr = st.element.get_or_add_pPr()
        bidi = pPr.find(qn('w:bidi'))
        if bidi is None:
            bidi = OxmlElement('w:bidi')
            pPr.append(bidi)
        rPr = st.element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:cs'), 'Vazirmatn')
        rFonts.set(qn('w:ascii'), 'Vazirmatn')
        rFonts.set(qn('w:hAnsi'), 'Vazirmatn')
        szCs = rPr.find(qn('w:szCs'))
        if szCs is None:
            szCs = OxmlElement('w:szCs')
            rPr.append(szCs)
        szCs.set(qn('w:val'), str(int(size * 2)))

    # --- Caption style (for figure/table captions) ---
    try:
        cap = doc.styles['Caption']
        cap.font.name = "Vazirmatn"
        cap.font.size = Pt(10)
        cap.font.italic = False
        cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr = cap.element.get_or_add_pPr()
        bidi = pPr.find(qn('w:bidi'))
        if bidi is None:
            pPr.append(OxmlElement('w:bidi'))
        rPr = cap.element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts')) or OxmlElement('w:rFonts')
        if rFonts.getparent() is None:
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:cs'), 'Vazirmatn')
    except KeyError:
        pass

    doc.save(str(REF_DOCX))
    print(f"Wrote {REF_DOCX}")


# ===================================================================
# 3) Run pandoc
# ===================================================================
def run_pandoc():
    cmd = [
        "pandoc",
        str(MD_TMP),
        "--from", "markdown",
        "--to", "docx",
        "--reference-doc", str(REF_DOCX),
        "--resource-path", str(ROOT / "figures_fa") + ";" + str(ROOT / "figures"),
        "-o", str(OUT_DOCX),
    ]
    print("Running:", " ".join(cmd))
    subprocess.check_call(cmd, cwd=str(HERE))
    print(f"Wrote {OUT_DOCX}")


# ===================================================================
# 4) Post-process the .docx to flip RTL on every paragraph and set fonts
# ===================================================================
def postprocess_docx():
    """Walk every paragraph, set RTL bidi + Vazirmatn complex-script font.
    English-block markers (<!--ENGLISH_BLOCK_START/END-->) toggle LTR + EB
    Garamond for the paragraphs between them.
    """
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document(str(OUT_DOCX))

    def set_rtl(paragraph):
        pPr = paragraph._p.get_or_add_pPr()
        if pPr.find(qn('w:bidi')) is None:
            pPr.append(OxmlElement('w:bidi'))

    def unset_rtl(paragraph):
        pPr = paragraph._p.get_or_add_pPr()
        bidi = pPr.find(qn('w:bidi'))
        if bidi is not None:
            pPr.remove(bidi)

    def set_run_fonts(run, *, latin="EB Garamond", cs="Vazirmatn"):
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), latin)
        rFonts.set(qn('w:hAnsi'), latin)
        rFonts.set(qn('w:cs'), cs)
        # Also set complex-script bold/italic mirroring main bold/italic
        if run.bold:
            bCs = rPr.find(qn('w:bCs'))
            if bCs is None:
                rPr.append(OxmlElement('w:bCs'))
        if run.italic:
            iCs = rPr.find(qn('w:iCs'))
            if iCs is None:
                rPr.append(OxmlElement('w:iCs'))

    in_english = False
    paragraphs_to_remove = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text == "<!--ENGLISH_BLOCK_START-->":
            in_english = True
            paragraphs_to_remove.append(para)
            continue
        if text == "<!--ENGLISH_BLOCK_END-->":
            in_english = False
            paragraphs_to_remove.append(para)
            continue
        if text == "<!--LTR_BLOCK_START-->":
            in_english = True
            paragraphs_to_remove.append(para)
            continue
        if text == "<!--LTR_BLOCK_END-->":
            in_english = False
            paragraphs_to_remove.append(para)
            continue

        if in_english:
            unset_rtl(para)
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                set_run_fonts(run, latin="EB Garamond", cs="EB Garamond")
        else:
            set_rtl(para)
            if para.style.name == 'Normal':
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in para.runs:
                # Vazirmatn for Persian/Arabic; Amiri may be better for Arabic
                # verses but full-string detection is brittle; we rely on
                # Vazirmatn falling back to Amiri for missing glyphs.
                set_run_fonts(run, latin="EB Garamond", cs="Vazirmatn")

    # Remove marker paragraphs
    for p in paragraphs_to_remove:
        p._element.getparent().remove(p._element)

    # Process tables: set RTL on every cell paragraph + Vazirmatn font
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    set_rtl(para)
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        set_run_fonts(run, latin="EB Garamond", cs="Vazirmatn")

    # Set document-wide RTL on the body
    body = doc.element.body
    sectPr = body.find(qn('w:sectPr'))
    if sectPr is not None:
        bidi = sectPr.find(qn('w:bidi'))
        if bidi is None:
            sectPr.append(OxmlElement('w:bidi'))

    doc.save(str(OUT_DOCX))
    print(f"Post-processed {OUT_DOCX} — RTL applied, fonts set, markers removed.")


# ===================================================================
# Entry point
# ===================================================================
def main():
    print("Step 1: Preprocessing markdown...")
    md_text = MD_IN.read_text(encoding='utf-8')
    md_text = preprocess_markdown(md_text)
    MD_TMP.write_text(md_text, encoding='utf-8')

    print("Step 2: Building reference.docx with Persian styles...")
    build_reference_docx()

    print("Step 3: Running pandoc...")
    run_pandoc()

    print("Step 4: Post-processing .docx for RTL + fonts...")
    postprocess_docx()

    # Cleanup
    MD_TMP.unlink(missing_ok=True)
    print("\nDone. Output:", OUT_DOCX)


if __name__ == "__main__":
    main()
